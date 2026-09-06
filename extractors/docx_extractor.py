import io
import os
import zipfile
from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    Safely extracts plain text and tables from a .docx file.
    Validates zip packaging and uses memory streams to prevent OS buffer locks.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found at: {file_path}")

    if os.path.getsize(file_path) == 0:
        raise ValueError("Uploaded file is empty (0 bytes).")

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    # Validate valid OpenXML ZIP archive
    if not zipfile.is_zipfile(io.BytesIO(file_bytes)):
        # Fallback: attempt plain text recovery if it is raw text or mislabeled
        try:
            raw_text = file_bytes.decode("utf-8", errors="ignore")
            clean = "".join(ch for ch in raw_text if ch.isprintable() or ch in "\n\t")
            if len(clean.strip()) > 40:
                return clean.strip()
        except Exception:
            pass
        raise ValueError(
            "Uploaded file is not a valid modern .docx package. "
            "If this file is an older Word (.doc) or RTF document, please save/export it as .docx or .pdf."
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
        extracted_text = []

        # 1. Paragraphs
        for paragraph in doc.paragraphs:
            cleaned_line = paragraph.text.strip()
            if cleaned_line:
                extracted_text.append(cleaned_line)

        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    extracted_text.append(" | ".join(row_data))

        result = "\n".join(extracted_text).strip()
        if not result:
            raise ValueError("No readable text found inside the Word document.")

        return result

    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Failed to process Word package: {str(e)}")