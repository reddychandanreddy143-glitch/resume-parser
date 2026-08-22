import pypdfium2 as pdfium
import io

# We will create a clean text PDF using standard python-docx conversion or direct write
from docx import Document

# Create a sample document first
doc = Document()
doc.add_heading('Chandan Reddy', level=1)
doc.add_paragraph('Email: chandan.reddy@example.com | Phone: +91 9876543210')

doc.add_heading('TECHNICAL SKILLS', level=2)
doc.add_paragraph('Python, Java, SQL, Flask, Git, PostgreSQL, NLP, Machine Learning')

doc.add_heading('EDUCATION', level=2)
doc.add_paragraph('Bachelor of Computer Applications (BCA)\nBengaluru North University (2023 - 2026)')

doc.add_heading('EXPERIENCE', level=2)
doc.add_paragraph('Python Developer Intern - Codec Technologies\n- Built NLP pipelines and Flask REST APIs.')

doc.add_heading('PROJECTS', level=2)
doc.add_paragraph('Automated Resume Parser: ATS document processing system using Python and Flask.')

# Save DOCX
doc.save('sample_resume.docx')
print("sample_resume.docx created successfully!")